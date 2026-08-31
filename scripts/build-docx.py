# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx", "Pillow"]
# ///
"""Build one book into a reading/printing DOCX, plates set at the moment they depict.

Usage:
    uv run scripts/build-docx.py                # -> build/Book-One.docx
    uv run scripts/build-docx.py book-two

Trim is 6x9 (author, s55). Each adopted plate in 08-Plates/images/plates/<book>/ is
placed after the paragraph its own sheet quotes from the page, so the picture
punctuates the beat rather than pre-empting it. Where a sheet is a craft record
with no page quote, the anchor is in MANUAL below — a snippet of the paragraph
the plate follows. build/ is derived but tracked (author, s55) — built at a
book's finish, not per session; build/plates-jpg/ is a cache and stays ignored.
"""
import re, os, sys, glob, subprocess, unicodedata
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BOOK = sys.argv[1] if len(sys.argv) > 1 else "book-one"
CHAPS = os.path.join(ROOT, "manuscript", BOOK)
PLATES = os.path.join(ROOT, "08-Plates/images/plates")
PORTRAITS_DIR = os.path.join(ROOT, "08-Plates/images/portraits")
OUTDIR = os.path.join(ROOT, "build")
JPGDIR = os.path.join(OUTDIR, "plates-jpg")

BODY_FONT, BODY_PT = "Garamond", 11.5
PAGE_W, PAGE_H = 6.0, 9.0
MARGIN = 0.7
TEXT_W = PAGE_W - 2 * MARGIN          # 4.6in
PLATE_MAX_W, PLATE_MAX_H = 4.6, 6.2
PORTRAIT_MAX_W, PORTRAIT_MAX_H = 2.9, 4.0   # a face reads smaller than a moment

# ---------------------------------------------------------------- anchors ---
# Plates whose sheets carry no usable page quote. Value: a snippet of the
# paragraph the plate should follow, "SAME:<slug>" to stack, or "LAST".
MANUAL = {
 (1,  "fallen-stars"):                   "and the doubled ones moved",
 (1,  "vask-kael-goat"):                 "because the goat was upset about the storm",
 (2,  "long-game"):                      "became the bottom of it",
 (3,  "the-wrist"):                      "and was proud of the quiet hands",
 (4,  "the-master-arrives"):             "the way kael's mother looked at the sea",
 (5,  "hand-on-the-rock"):               "cold under his hand, carrying nothing",
 (5,  "past-the-split-rock"):            "he stepped past the rock",
 (6,  "carried-past-the-rock"):          "the silence was full of the third voice",
 (7,  "the-wasters"):                    "watched him find the weight of it",
 (8,  "the-one-who-did-not-laugh-cold"): "SAME:the-one-who-did-not-laugh",
 (12, "the-law-on-the-table"):           "looked at the law lying on the table",
 (21, "dead-restated"):                  "over the heart, and stayed there",
 (24, "turn-him-toward-the-lamp"):       '"turn him toward the lamp," the princess said',
 (25, "both-hands"):                     "she said nothing. then she let go",
 (26, "the-fan-downrange"):              "SAME:the-fan",
 (28, "the-hand-back"):                  "open, not touching anything",
 (32, "the-white-in-her-hair"):          "the white of the seam in a stone when you split it",
 (33, "the-lamp"):                       "then she came back, and picked the lamp up",
 (34, "the-bowl"):                       "they came out of the other side of it still going",
 (35, "the-white-room"):                 "going up and going out and going up again",
 (35, "the-spotter"):                    "near the part of kael that was working",
 (36, "no-lights"):                      "LAST",
}

# -------------------------------------------------------------- portraits ---
# References set at the moment the book first puts that person in front of the
# reader. `images/portraits/` are references, not pages (08-Plates/README.md
# §The files) — placing them here makes them pages, which is the author's call
# (s55). Book One ages only: kael-17, valeria-17 and aeliana-18 are Book Two+
# states and are left out. (chapter, block the portrait follows, file, why)
PORTRAITS = [
    ( 1,  1, "kael-4.png",     "the boy in the loft the epigraph names"),
    ( 1,  9, "vask.png",       "the naming - a stone put down on another stone"),
    ( 1, 11, "neris.png",      "counting the boats; the stillness is the portrait"),
    ( 4, 12, "severin.png",    '"This is the master," his mother said'),
    ( 5,  4, "kael-9.png",     "the new school, the year he turns nine"),
    (13, 26, "elarine-14.png", '"Again," Elarine said. "Smaller."'),
    (13, 55, "aurelian-14.png","stepped into the cistern court"),
    (14, 97, "valeria-14.png", "a girl his own age had just walked through it"),
    (22, 60, "kael-14.png",    "he was thirteen; since he was four, to be seen"),
    (24, 27, "aeliana-15.png", "her hair was down - Ch.22 wears it up, the reference is loose"),
]

def norm(s):
    s = unicodedata.normalize("NFKC", s)
    for a, b in (("’","'"),("‘","'"),("“",'"'),("”",'"'),
                 ("—","-"),("–","-"),("−","-")):
        s = s.replace(a, b)
    return re.sub(r"\s+", " ", re.sub(r"[*_`]", "", s)).strip().lower()

def sheet_sections(t):
    """The CHOSEN bullets of a plate-candidates file, keyed by plate slug.

    Since s58 the chapter's candidate list at 08-Plates/prompts/plate-candidates/<book>/
    replaces the retired 08-Plates/plates/ sheets. A CHOSEN bullet names its plate
    by the prompt path it points at, and carries the chapter quote that anchors the
    picture (see canon_frags). The old sheets are in 08-Plates/archive/plate-sheets.md.
    """
    out = {}
    m = re.search(r"^## CHOSEN\s*$(.*?)(?=^## |\Z)", t, flags=re.M | re.S)
    for line in (m.group(1).split("\n") if m else []):
        s = re.search(r"`[^`]*plates/(?:book-[a-z]+/)?ch\d+-([a-z0-9-]+)\.md`", line)
        if s:
            out[s.group(1)] = line
    return out

def canon_frags(body):
    body = re.sub(r"```.*?```", "", body, flags=re.S)
    m = re.search(r"canon:(.*?)(?:\n\n|\Z)", body, flags=re.S)
    scope = m.group(1) if m else body
    frags = []
    for sp in re.findall(r"(?<!\*)\*([^*\n][^*]*)\*(?!\*)", scope):
        if len(sp) <= 25:
            continue
        for f in re.split(r"\s*(?:…|\.\.\.|\s/\s)\s*", sp):
            if len(norm(f)) > 25:
                frags.append(norm(f))
    return frags

def resolve_images(chapters):
    adopted = {}
    for p in sorted(glob.glob(os.path.join(PLATES, "*", "*.png"))):
        m = re.match(r"ch(\d+)-(.+)$", os.path.basename(p)[:-4])
        if m and int(m.group(1)) in chapters:
            adopted.setdefault(int(m.group(1)), []).append((m.group(2), p))
    placed, missed = {}, []
    for n in sorted(adopted):
        hits = glob.glob(os.path.join(ROOT, "08-Plates/prompts/plate-candidates/*/ch%02d.md" % n))
        secs = sheet_sections(open(hits[0]).read()) if hits else {}
        nb = chapters[n]["norm"]
        for slug, path in adopted[n]:
            idx = None
            if (n, slug) in MANUAL:
                want = MANUAL[(n, slug)]
                if want == "LAST":
                    idx = len(nb) - 1
                elif want.startswith("SAME:"):
                    idx = ("SAME", want[5:])
                else:
                    hits = [i for i, b in enumerate(nb) if want in b]
                    idx = hits[0] if hits else None
            else:
                body = secs.get(slug) or next(
                    (v for k, v in secs.items() if slug.startswith(k) or k.startswith(slug)), "")
                hits = [i for f in canon_frags(body)
                          for i, b in enumerate(nb) if f in b]
                idx = max(hits) if hits else None
            if idx is None:
                missed.append("ch%02d-%s" % (n, slug))
            else:
                placed[(n, slug)] = [idx, path]
    for key, v in placed.items():
        if isinstance(v[0], tuple):
            v[0] = placed[(key[0], v[0][1])][0]
    if missed:
        sys.exit("unanchored plates (add to MANUAL): " + ", ".join(missed))
    by_chapter = {}
    for (n, slug), (idx, path) in placed.items():
        by_chapter.setdefault(n, {}).setdefault(idx, []).append((slug, path, "plate"))
    for n in by_chapter:
        for idx in by_chapter[n]:
            by_chapter[n][idx].sort()
    n_port = 0
    for chap, blk, fn, _why in PORTRAITS:
        src = os.path.join(PORTRAITS_DIR, fn)
        if not os.path.exists(src):
            sys.exit("missing portrait: " + src)
        if chap not in chapters or blk >= len(chapters[chap]["blocks"]):
            sys.exit("portrait anchor out of range: %s ch%02d block %d" % (fn, chap, blk))
        by_chapter.setdefault(chap, {}).setdefault(blk, []).append((fn, src, "portrait"))
        n_port += 1
    return by_chapter, len(placed), n_port

# ------------------------------------------------------------------ images --
def jpeg_for(png):
    os.makedirs(JPGDIR, exist_ok=True)
    jpg = os.path.join(JPGDIR, os.path.basename(png)[:-4] + ".jpg")
    if not os.path.exists(jpg) or os.path.getmtime(jpg) < os.path.getmtime(png):
        im = Image.open(png).convert("RGB")
        if im.width > 1400:
            im = im.resize((1400, round(im.height * 1400 / im.width)), Image.LANCZOS)
        im.save(jpg, "JPEG", quality=88, optimize=True, progressive=True)
    return jpg

# -------------------------------------------------------------- docx parts --
def field(par, code):
    r = par.add_run()
    for tag, attr, txt in (("w:fldChar", "begin", None),
                           ("w:instrText", None, code),
                           ("w:fldChar", "end", None)):
        el = OxmlElement(tag)
        if attr:
            el.set(qn("w:fldCharType"), attr)
        if txt is not None:
            el.set(qn("xml:space"), "preserve"); el.text = txt
        r._r.append(el)

def style_page(section, numbered):
    section.page_width, section.page_height = Inches(PAGE_W), Inches(PAGE_H)
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(MARGIN)
    section.footer_distance = Inches(0.45)
    if numbered:
        pgnum = OxmlElement("w:pgNumType"); pgnum.set(qn("w:start"), "1")
        section._sectPr.append(pgnum)
        section.footer.is_linked_to_previous = False
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        field(p, "PAGE")
        for r in p.runs:
            r.font.name, r.font.size = BODY_FONT, Pt(9.5)
    else:
        section.footer.is_linked_to_previous = False
        section.footer.paragraphs[0].text = ""

def para(doc, text="", *, size=BODY_PT, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         indent=0.0, before=0, after=0, italic=False, bold=False, caps=False,
         spacing=1.22, keep=False, color=None, inset=0.0, md=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.first_line_indent = Inches(indent)
    pf.left_indent = pf.right_indent = Inches(inset)
    pf.space_before, pf.space_after = Pt(before), Pt(after)
    pf.line_spacing = spacing
    pf.widow_control = True
    pf.keep_with_next = keep
    runs = md_runs(p, text) if md else [p.add_run(text)]
    for r in runs:
        f = r.font
        f.name, f.size = BODY_FONT, Pt(size)
        if italic: f.italic = True
        if bold:   f.bold = True
        if caps:   f.small_caps = True
        if color:  f.color.rgb = color
        r._element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)
    return p

def md_runs(p, text):
    out = []
    for tok in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.italic = True; r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            r = p.add_run(tok)
        out.append(r)
    return out

def add_image(doc, png, kind="plate"):
    jpg = jpeg_for(png)
    with Image.open(jpg) as im:
        w, h = im.size
    mw, mh = ((PLATE_MAX_W, PLATE_MAX_H) if kind == "plate"
              else (PORTRAIT_MAX_W, PORTRAIT_MAX_H))
    width = min(mw, mh * w / h)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before, pf.space_after = Pt(14), Pt(14)
    pf.first_line_indent = Inches(0)
    pf.keep_together = True
    p.add_run().add_picture(jpg, width=Inches(width))

# ------------------------------------------------------------------- build --
def main():
    if not os.path.isdir(CHAPS):
        sys.exit("no such book: " + CHAPS)
    chapters = {}
    for f in sorted(glob.glob(os.path.join(CHAPS, "*.md"))):
        n = int(os.path.basename(f)[:2])
        blocks = [b.strip() for b in re.split(r"\n\s*\n", open(f).read()) if b.strip()]
        chapters[n] = {"blocks": blocks, "norm": [norm(b) for b in blocks]}
    plates, n_plates, n_portraits = resolve_images(chapters)

    title = re.sub(r"-", " ", BOOK).title()
    try:
        sha = subprocess.check_output(["git", "-C", ROOT, "rev-parse", "--short", "HEAD"]).decode().strip()
        date = subprocess.check_output(["git", "-C", ROOT, "log", "-1", "--format=%cs"]).decode().strip()
    except Exception:
        sha = date = "?"

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name, st.font.size = BODY_FONT, Pt(BODY_PT)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    hy = OxmlElement("w:autoHyphenation"); hy.set(qn("w:val"), "true")
    doc.settings.element.append(hy)
    doc.core_properties.title = title
    doc.core_properties.author = "Sebas"

    style_page(doc.sections[0], numbered=False)

    # -- title page
    for _ in range(6):
        para(doc, "", after=0)
    para(doc, title.upper(), size=30, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    para(doc, "an epic fantasy", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=64)
    para(doc, "Sebas", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, caps=True)

    # -- colophon
    doc.add_page_break()
    for _ in range(14):
        para(doc, "", after=0)
    grey = RGBColor(0x44, 0x44, 0x44)
    for line in (
        "Draft manuscript. Working title to be decided.",
        "",
        "%d chapters, %d plates, %d portraits." % (len(chapters), n_plates, n_portraits),
        "Each plate is set after the passage it depicts.",
        "",
        "Built from the repository at commit %s (%s)." % (sha, date),
        "",
        "This work is licensed under a Creative Commons",
        "Attribution-ShareAlike 4.0 International License.",
    ):
        para(doc, line, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=grey, after=2)

    # -- contents
    doc.add_page_break()
    para(doc, "Contents", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, caps=True, after=22)
    for n in sorted(chapters):
        head = chapters[n]["blocks"][0].lstrip("# ").strip()
        para(doc, head, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=7)

    # -- body
    body = doc.add_section(WD_SECTION.ODD_PAGE)
    style_page(body, numbered=True)

    for ci, n in enumerate(sorted(chapters)):
        blocks = chapters[n]["blocks"]
        marks = plates.get(n, {})
        if ci:
            doc.add_page_break()
        head = blocks[0].lstrip("# ").strip()
        num, _, name = head.partition("—")
        for _ in range(3):
            para(doc, "", after=0)
        para(doc, num.strip(), size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER,
             caps=True, after=8, keep=True)
        para(doc, name.strip(), size=17, align=WD_ALIGN_PARAGRAPH.CENTER,
             after=26, keep=True)
        for _slug, path, kind in marks.get(0, []):
            add_image(doc, path, kind)

        fresh = True   # next paragraph opens a scene: no first-line indent
        for bi in range(1, len(blocks)):
            b = blocks[bi]
            if re.fullmatch(r"-{3,}", b):
                para(doc, " * * * ", align=WD_ALIGN_PARAGRAPH.CENTER,
                     before=12, after=12)
                fresh = True
            else:
                text = re.sub(r"\s*\n\s*", " ", b)
                is_epigraph = bi == 1 and text.startswith("*") and text.endswith("*")
                para(doc, text, md=True, indent=0 if (fresh or is_epigraph) else 0.2,
                     inset=0.25 if is_epigraph else 0.0,
                     after=10 if is_epigraph else 0)
                fresh = False
            for _slug, path, kind in marks.get(bi, []):
                add_image(doc, path, kind)
                fresh = True

    os.makedirs(OUTDIR, exist_ok=True)
    out = os.path.join(OUTDIR, title.replace(" ", "-") + ".docx")
    doc.save(out)
    print("built: %s (%.1f MB, %d chapters, %d plates, %d portraits)"
          % (out, os.path.getsize(out) / 1e6, len(chapters), n_plates, n_portraits))

main()
